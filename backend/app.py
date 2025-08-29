from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import uvicorn
import os
import json
import csv
import io
import uuid
import tempfile
from datetime import datetime
from typing import List, Optional
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import aiofiles

from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from azure.cognitiveservices.vision.computervision.models import OperationStatusCodes
from msrest.authentication import CognitiveServicesCredentials
from dotenv import load_dotenv
import requests
from PIL import Image

# Try to import python-docx for DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX files will not be processed.")

# Load environment variables
load_dotenv()
AZURE_API_KEY = os.getenv('AZURE_API_KEY')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
AZURE_ENDPOINT = os.getenv('AZURE_ENDPOINT', 'https://your-resource-name.cognitiveservices.azure.com/')

# Create FastAPI app
app = FastAPI(title="SoF Event Extractor API", version="2.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories for storing jobs and results
JOBS_DIR = Path("jobs")
RESULTS_DIR = Path("results")
JOBS_DIR.mkdir(exist_ok=True)
RESULTS_DIR.mkdir(exist_ok=True)

# Job storage (in production, use a database)
jobs = {}

# Thread pool for CPU-bound tasks
executor = ThreadPoolExecutor(max_workers=4)

def save_job_status(job_id: str, status: str, progress: int = 0, message: str = "", results: dict = None, total_files: int = 1, filenames: list = None):
    """Save job status to file"""
    job_data = {
        "job_id": job_id,
        "status": status,
        "progress": progress,
        "message": message,
        "created_at": datetime.now().isoformat(),
        "total_files": total_files,
        "filenames": filenames or [],
        "results": results or {}
    }

    with open(JOBS_DIR / f"{job_id}.json", "w") as f:
        json.dump(job_data, f, indent=2)

    jobs[job_id] = job_data

def load_job_status(job_id: str) -> Optional[dict]:
    """Load job status from file"""
    job_file = JOBS_DIR / f"{job_id}.json"
    if job_file.exists():
        try:
            with open(job_file, "r") as f:
                return json.load(f)
        except:
            return None
    return None

def regenerate_csv_if_needed(job_id: str, results_data: dict):
    """Regenerate CSV file if it contains invalid data"""
    csv_path = RESULTS_DIR / f"{job_id}.csv"

    # Check if CSV exists and contains valid CSV data
    if csv_path.exists():
        with open(csv_path, "r", encoding='utf-8') as f:
            content = f.read().strip()
            # If it starts with {, it's JSON, not CSV
            if content.startswith('{'):
                print(f"CSV file {job_id}.csv contains JSON data, regenerating...")
                # Regenerate the CSV
                all_rows = results_data.get("table", [])
                if all_rows:
                    # internal helper to extract a YYYY-MM-DD day value from a row
                    def date_from_row_local(r):
                        d = r.get('day') or ''
                        if d:
                            if isinstance(d, str) and len(d) >= 8 and '-' in d:
                                return d
                            try:
                                dt = datetime.fromisoformat(str(d).replace(' ', 'T'))
                                return dt.strftime('%Y-%m-%d')
                            except Exception:
                                return d

                        s = r.get('start_time') or r.get('start') or ''
                        if s:
                            try:
                                dt = datetime.fromisoformat(str(s).replace(' ', 'T'))
                                return dt.strftime('%Y-%m-%d')
                            except Exception:
                                try:
                                    return str(s).split(' ')[0]
                                except Exception:
                                    return ''

                        return ''

                    csv_buffer = io.StringIO()
                    fieldnames = ['event', 'day', 'start', 'end', 'duration', 'ship_cargo', 'layoff_time', 'description', 'filename']
                    writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames, quoting=csv.QUOTE_ALL, escapechar='\\')
                    writer.writeheader()

                    for row in all_rows:
                        # ensure defaults and sanitize
                        duration_val = row.get('duration', '')
                        if isinstance(duration_val, str):
                            duration_val = ' '.join(duration_val.split())

                        ship = row.get('ship_cargo') or row.get('shipCargo') or 'N/A'
                        lay = row.get('layoff_time') or row.get('layoff') or ''
                        if not lay:
                            lay = 'N/A'

                        csv_row = {
                            'event': row.get('event', ''),
                            'day': date_from_row_local(row),
                            'start': row.get('start_time', '') or row.get('start', ''),
                            'end': row.get('end_time', '') or row.get('end', ''),
                            'duration': duration_val,
                            'ship_cargo': ship,
                            'layoff_time': lay,
                            'description': row.get('description', ''),
                            'filename': row.get('filename', '')
                        }
                        writer.writerow(csv_row)

                    with open(csv_path, "w", newline='', encoding='utf-8') as f:
                        f.write(csv_buffer.getvalue())
                    print(f"CSV file regenerated: {csv_path}")

def process_with_gemini(text: str, filename: str) -> List[dict]:
    """Process extracted text with Gemini to extract structured data."""
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"

        prompt = f"""
        Extract event data from this OCR text and return as JSON array.

        Text: {text}

        IMPORTANT INSTRUCTIONS:
        1. Calculate DURATION: If you have start_time and end_time, calculate the duration in hours and minutes (e.g., "2h 30m", "1h 15m", "45m")
        2. For durations ABOVE 24 hours: Use days and hours format (e.g., "2d 4h 30m", "1d 12h", "3d 2h 15m")
        3. Make EVENT names descriptive and professional (e.g., "Cargo Loading Operation", "Ship Maintenance Activity", "Crew Briefing Session")
        4. Make DESCRIPTION more detailed and meaningful, explaining what actually happened during the event
        5. Use proper date formats: YYYY-MM-DD HH:MM for start_time and end_time
        6. Extract all relevant information about ships, cargo, layoff times, etc.
        7. For LAYOFF_TIME: Look for any time periods that could be layoff, break, rest, or pause times. Format it the same way as duration (e.g., "2h 0m", "30m"). If you find any time information that might be layoff time, extract it. Only use "N/A" if there's absolutely no time information that could be related to layoff/break periods.
        8. For SHIP_CARGO: If no ship or cargo information is available, use "N/A"

        Return ONLY a JSON array of objects with these exact fields:
        - event: descriptive event name (not just short codes)
        - day: day of week
        - start_time: start time with date (YYYY-MM-DD HH:MM format)
        - end_time: end time with date (YYYY-MM-DD HH:MM format)
        - duration: calculated duration from start and end times (e.g., "2h 30m" or "2d 4h 30m" for longer durations)
        - ship_cargo: ship/cargo information or "N/A" if not available
        - layoff_time: any layoff/break/rest time period in duration format (e.g., "2h 0m", "30m") or "N/A" if none found
        - description: detailed description of what happened
        - filename: "{filename}"

        If no events found, return empty array [].
        """

        payload = {
            "contents": [{
                "parts": [{"text": prompt}]
            }]
        }

        response = requests.post(url, json=payload, headers={"Content-Type": "application/json"})
        response.raise_for_status()

        result = response.json()
        generated_text = result['candidates'][0]['content']['parts'][0]['text']

        # Extract JSON from response
        start = generated_text.find('[')
        end = generated_text.rfind(']') + 1
        if start != -1 and end != -1:
            json_str = generated_text[start:end]
            return json.loads(json_str)
        else:
            return []

    except Exception as e:
        print(f"Gemini processing failed for {filename}: {str(e)}")
        return []

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install with: pip install python-docx")

    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        raise

async def process_files_background(job_id: str, files: List[UploadFile], use_enhanced_processing: bool = False):
    """Background task to process uploaded files"""
    try:
        print(f"DOCX_AVAILABLE: {DOCX_AVAILABLE}")
        save_job_status(job_id, "processing", 10, "Initializing Azure OCR client...")

        # Initialize Azure client
        computervision_client = ComputerVisionClient(
            AZURE_ENDPOINT,
            CognitiveServicesCredentials(AZURE_API_KEY)
        )

        all_rows = []
        total_files = len(files)

        for i, file in enumerate(files):
            try:
                progress = 10 + (i / total_files) * 80
                save_job_status(job_id, "processing", int(progress), f"Processing {file.filename}...")

                # Save uploaded file to temp
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
                    content = await file.read()
                    tmp.write(content)
                    temp_path = tmp.name

                try:
                    # Check if this is a DOCX file
                    is_docx = (file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
                              file.filename.lower().endswith('.docx'))

                    print(f"Processing file: {file.filename}, content_type: {file.content_type}, is_docx: {is_docx}")

                    if is_docx:
                        # Process DOCX file
                        if not DOCX_AVAILABLE:
                            print(f"Skipping {file.filename}: python-docx not installed")
                            continue

                        try:
                            extracted_text = extract_text_from_docx(temp_path)
                            print(f"Extracted {len(extracted_text)} characters from DOCX: {file.filename}")
                            print(f"First 200 chars: {extracted_text[:200]}")
                        except Exception as e:
                            print(f"Error extracting text from DOCX {file.filename}: {str(e)}")
                            continue
                    else:
                        # Process image file with Azure OCR
                        print(f"Processing as image: {file.filename}")
                        with open(temp_path, "rb") as image_file:
                            # Call Azure OCR
                            ocr_result = computervision_client.read_in_stream(image_file, raw=True)
                        operation_id = ocr_result.headers['Operation-Location'].split('/')[-1]

                        # Wait for result
                        while True:
                            result = computervision_client.get_read_result(operation_id)
                            if result.status not in [OperationStatusCodes.running, OperationStatusCodes.not_started]:
                                break
                            await asyncio.sleep(1)

                        extracted_text = ""
                        if result.status == OperationStatusCodes.succeeded:
                            for page in result.analyze_result.read_results:
                                for line in page.lines:
                                    extracted_text += line.text + "\n"

                    # Process with Gemini if text was extracted
                    if extracted_text.strip():
                        print(f"Processing {len(extracted_text)} characters with Gemini for {file.filename}")
                        gemini_result = await asyncio.get_event_loop().run_in_executor(
                            executor, process_with_gemini, extracted_text, file.filename
                        )
                        if gemini_result:
                            all_rows.extend(gemini_result)
                            print(f"Gemini returned {len(gemini_result)} events for {file.filename}")
                        else:
                            print(f"Gemini returned no events for {file.filename}")
                    else:
                        print(f"No text extracted from {file.filename}")

                finally:
                    # Cleanup temp file
                    try:
                        os.unlink(temp_path)
                    except:
                        pass

            except Exception as e:
                print(f"Error processing {file.filename}: {str(e)}")
                continue

        # Save results
        if all_rows:
            results_data = {"table": all_rows}

            # Save JSON results
            with open(RESULTS_DIR / f"{job_id}.json", "w") as f:
                json.dump(results_data, f, indent=2)

            # Save CSV results with frontend-compatible field names
            csv_buffer = io.StringIO()
            if all_rows:
                # Use the same field names that are sent to frontend
                fieldnames = ['event', 'day', 'start_time', 'end_time', 'duration', 'ship_cargo', 'layoff_time', 'description', 'filename']
                writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames, quoting=csv.QUOTE_ALL, escapechar='\\')
                writer.writeheader()

                for row in all_rows:
                    # Map backend field names to frontend field names for CSV export
                    csv_row = {
                        'event': row.get('event', ''),
                        'day': row.get('day', ''),
                        'start_time': row.get('start_time', ''),  # Keep full field name
                        'end_time': row.get('end_time', ''),      # Keep full field name
                        'duration': row.get('duration', ''),
                        'ship_cargo': row.get('ship_cargo', ''),
                        'layoff_time': row.get('layoff_time', ''),
                        'description': row.get('description', ''),
                        'filename': row.get('filename', '')
                    }
                    writer.writerow(csv_row)

                # Get the CSV content
                csv_content = csv_buffer.getvalue()
                print(f"Generated CSV content length: {len(csv_content)}")  # Debug log

                # Write CSV to file
                with open(RESULTS_DIR / f"{job_id}.csv", "w", newline='', encoding='utf-8') as f:
                    f.write(csv_content)

                print(f"CSV file saved: {RESULTS_DIR / f'{job_id}.csv'}")  # Debug log

            save_job_status(job_id, "completed", 100, f"Successfully processed {len(all_rows)} rows from {total_files} files", results_data)
            print(f"Job {job_id} completed successfully with {len(all_rows)} events")
        else:
            save_job_status(job_id, "completed", 100, f"Processed {total_files} files but no structured data found")
            print(f"Job {job_id} completed but no events found")

    except Exception as e:
        error_msg = str(e)
        if 'your-resource-name' in error_msg:
            save_job_status(job_id, "failed", 0, "Azure endpoint configuration error. Please update AZURE_ENDPOINT in .env file.")
        elif 'authentication' in error_msg.lower() or 'unauthorized' in error_msg.lower():
            save_job_status(job_id, "failed", 0, "Azure API key authentication failed. Please check AZURE_API_KEY in .env file.")
        else:
            save_job_status(job_id, "failed", 0, f"Processing error: {error_msg}")

@app.get("/")
async def root():
    """Health check endpoint"""
    return {"message": "SoF Event Extractor API", "version": "2.0.0", "status": "running"}

@app.post("/api/upload")
async def upload_files(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    use_enhanced_processing: bool = False
):
    """Upload files for processing"""
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 files allowed")

    # Validate file types
    allowed_types = ['image/jpeg', 'image/png', 'image/jpg', 'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    for file in files:
        if file.content_type not in allowed_types and not file.content_type.startswith('image/'):
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")

    # Check file sizes (10MB limit per file)
    for file in files:
        if file.size and file.size > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"File too large: {file.filename}. Max size is 10MB per file.")

    # Generate job ID
    job_id = str(uuid.uuid4())

    # Extract filenames for metadata
    filenames = [file.filename for file in files]

    # Save initial job status
    save_job_status(job_id, "queued", 0, "Files uploaded, processing started", total_files=len(files), filenames=filenames)

    # Start background processing
    background_tasks.add_task(process_files_background, job_id, files, use_enhanced_processing)

    return {
        "job_id": job_id,
        "total_files": len(files),
        "message": "Files uploaded successfully. Processing started."
    }

@app.get("/api/jobs")
async def get_jobs():
    """Get all jobs"""
    job_files = list(JOBS_DIR.glob("*.json"))
    jobs_list = []

    for job_file in job_files:
        try:
            with open(job_file, "r") as f:
                job_data = json.load(f)
                jobs_list.append({
                    "job_id": job_data["job_id"],
                    "status": job_data["status"],
                    "progress": job_data.get("progress", 0),
                    "message": job_data.get("message", ""),
                    "created_at": job_data.get("created_at", ""),
                    "total_files": job_data.get("total_files", 0),
                    "filename": job_data.get("filenames", [None])[0] if job_data.get("filenames") else "",
                    "filenames": job_data.get("filenames", [])
                })
        except:
            continue

    # Sort by creation date (newest first)
    jobs_list.sort(key=lambda x: x["created_at"], reverse=True)

    return {"jobs": jobs_list}

@app.get("/api/result/{job_id}")
async def get_result(job_id: str):
    """Get results for a specific job"""
    job_status = load_job_status(job_id)
    if not job_status:
        raise HTTPException(status_code=404, detail="Job not found")

    if job_status["status"] == "processing":
        return {
            "job_id": job_id,
            "status": "processing",
            "progress": job_status.get("progress", 0),
            "message": job_status.get("message", "Processing in progress..."),
            "total_files": job_status.get("total_files", 1),
            "filename": job_status.get("filenames", [None])[0] if job_status.get("filenames") else "",
            "filenames": job_status.get("filenames", [])
        }

    if job_status["status"] == "failed":
        return {
            "job_id": job_id,
            "status": "failed",
            "message": job_status.get("message", "Processing failed"),
            "total_files": job_status.get("total_files", 1),
            "filename": job_status.get("filenames", [None])[0] if job_status.get("filenames") else "",
            "filenames": job_status.get("filenames", []),
            "error": job_status.get("message", "Processing failed")
        }

    # Load results
    results_file = RESULTS_DIR / f"{job_id}.json"
    if not results_file.exists():
        raise HTTPException(status_code=404, detail="Results not found")

    with open(results_file, "r") as f:
        results = json.load(f)

    # Regenerate CSV if needed to ensure proper format
    regenerate_csv_if_needed(job_id, results)

    # Return data in the format expected by the frontend
    # Map backend field names to frontend expected field names
    mapped_events = []
    for event in results.get("table", []):
        mapped_event = {
            "event": event.get("event", ""),
            "start": event.get("start_time", ""),
            "end": event.get("end_time", ""),
            "description": event.get("description", ""),
            "filename": event.get("filename", ""),
            "day": event.get("day", ""),
            "duration": event.get("duration", ""),
            "ship_cargo": event.get("ship_cargo", ""),
            "layoff_time": event.get("layoff_time", "")
        }
        mapped_events.append(mapped_event)

    return {
        "job_id": job_id,
        "status": "completed",
        "events": mapped_events,  # Frontend expects events array directly
        "total_files": job_status.get("total_files", 1),
        "filename": job_status.get("filenames", [None])[0] if job_status.get("filenames") else "",  # Get first filename from list
        "filenames": job_status.get("filenames", []),
        "message": job_status.get("message", ""),
        "created_at": job_status.get("created_at", "")
    }

@app.post("/api/export/{job_id}")
async def export_result(job_id: str, format: str = "json", type: Optional[str] = None, request: Request = None):
    """Export results in specified format.

    For CSV exports we now generate CSV on-the-fly from the JSON results to
    guarantee a valid CSV (avoids serving stale or accidentally-written JSON
    files with a .csv extension).
    """
    job_status = load_job_status(job_id)
    if not job_status or job_status["status"] != "completed":
        raise HTTPException(status_code=404, detail="Job not found or not completed")


    # Determine requested format - frontend may send ?type=csv
    out_format = type if type else format

    results_file = RESULTS_DIR / f"{job_id}.json"
    if not results_file.exists():
        raise HTTPException(status_code=404, detail="Results file not found")

    # Load JSON results once (fallback source)
    with open(results_file, "r", encoding='utf-8') as f:
        results = json.load(f)

    # If frontend POSTed a normalization payload (events), prefer that for export
    posted_events = None
    if request is not None:
        try:
            body = await request.json()
            if isinstance(body, dict) and 'events' in body:
                posted_events = body['events']
        except Exception:
            posted_events = None

    if out_format.lower() == "json":
        # Use edited data if posted, otherwise use stored results
        if posted_events:
            # Map frontend field names to backend field names for consistency
            mapped_events = []
            for row in posted_events:
                mapped_event = {
                    'event': row.get('event') or row.get('Event') or row.get('name') or '',
                    'day': row.get('day') or row.get('Day') or row.get('date') or '',
                    'start_time': row.get('start_time') or row.get('Start Time') or row.get('start_time_iso') or '',
                    'end_time': row.get('end_time') or row.get('End Time') or row.get('end_time_iso') or '',
                    'duration': row.get('duration') or row.get('Duration') or '',
                    'ship_cargo': row.get('ship_cargo') or row.get('Ship/Cargo') or row.get('ShipCargo') or 'N/A',
                    'layoff_time': row.get('layoff_time') or row.get('Layoff Time') or row.get('laytime') or row.get('Laytime') or 'N/A',
                    'description': row.get('description') or row.get('Description') or '',
                    'filename': row.get('filename') or row.get('Filename') or row.get('FileName') or ''
                }
                mapped_events.append(mapped_event)

            # Return the mapped events as JSON with backend field names
            return JSONResponse(content={"table": mapped_events})
        else:
            # Return stored results
            return FileResponse(
                results_file,
                media_type="application/json",
                filename=f"extracted_data_{job_id}.json"
            )


    elif out_format.lower() == "csv":
        # Use edited data if posted, otherwise use stored results
        events = posted_events if posted_events else (results.get("table") or results.get("events") or [])
        if not events:
            raise HTTPException(status_code=404, detail="No events found for this job.")

        # Always use the original backend field names for consistency
        # Map frontend field names to backend field names
        fieldnames = ['event', 'day', 'start_time', 'end_time', 'duration', 'ship_cargo', 'layoff_time', 'description', 'filename']

        csv_buffer = io.StringIO()
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames, quoting=csv.QUOTE_ALL, escapechar='\\')
        writer.writeheader()

        for row in events:
            # Map frontend field names to backend field names
            csv_row = {
                'event': row.get('event') or row.get('Event') or row.get('name') or '',
                'day': row.get('day') or row.get('Day') or row.get('date') or '',
                'start_time': row.get('start_time') or row.get('Start Time') or row.get('start_time_iso') or '',
                'end_time': row.get('end_time') or row.get('End Time') or row.get('end_time_iso') or '',
                'duration': row.get('duration') or row.get('Duration') or '',
                'ship_cargo': row.get('ship_cargo') or row.get('Ship/Cargo') or row.get('ShipCargo') or 'N/A',
                'layoff_time': row.get('layoff_time') or row.get('Layoff Time') or row.get('laytime') or row.get('Laytime') or 'N/A',
                'description': row.get('description') or row.get('Description') or '',
                'filename': row.get('filename') or row.get('Filename') or row.get('FileName') or ''
            }
            writer.writerow(csv_row)

        csv_bytes = csv_buffer.getvalue().encode('utf-8')
        csv_buffer.close()
        return StreamingResponse(
            iter([csv_bytes]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=extracted_data_{job_id}.csv"}
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported export format. Use 'json' or 'csv'")

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "azure_configured": bool(AZURE_API_KEY and AZURE_ENDPOINT and 'your-resource-name' not in AZURE_ENDPOINT),
        "gemini_configured": bool(GEMINI_API_KEY)
    }

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
